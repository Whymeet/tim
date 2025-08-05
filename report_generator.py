from collections import defaultdict
from docx import Document
from docx.shared import Inches
import os

def generate_report(posts: list[dict],
                    output_file: str = "Отчёт.docx",
                    assets_dir: str = "assets",
                    inner_image: str = "inner.png") -> None:
    doc = Document()
    doc.add_heading("Отчёт по рекламным кампаниям VK", level=0)

    # группируем посты по компании
    grouped = defaultdict(list)
    for p in posts:
        grouped[p["Компания"]].append(p)

    for company, items in grouped.items():
        doc.add_heading(company, level=1)

        for idx, post in enumerate(items, start=1):
            subtitle = f"{idx}. {post['Группа']}"
            doc.add_heading(subtitle, level=2)
            doc.add_paragraph(post["Ссылка"])

            # скрин поста
            if os.path.exists(post.get("Скриншот", "")):
                doc.add_picture(post["Скриншот"], width=Inches(5))

            # сначала добавляем overview_funnel сразу после поста
            prefix = post["Группа"].lower()
            funnel_file = f"{post['Группа']}_overview_funnel.png"
            funnel_path = os.path.join(assets_dir, funnel_file)
            if os.path.exists(funnel_path):
                doc.add_picture(funnel_path, width=Inches(5))
                print(f"✅ Добавлен overview_funnel для группы: {post['Группа']}")

            # затем все остальные скрины статистики (кроме funnel)
            stats = sorted(
                f for f in os.listdir(assets_dir)
                if f.lower().startswith(prefix) and f.lower().endswith(".png")
                and not f.endswith("_overview_funnel.png")  # исключаем funnel
            )
            for fname in stats:
                doc.add_picture(os.path.join(assets_dir, fname), width=Inches(5))

            # добавляем inner.png в конце каждой рекламной кампании
            if os.path.exists(inner_image):
                doc.add_picture(inner_image, width=Inches(5))
                print(f"✅ Добавлен {inner_image} для группы: {post['Группа']}")
            else:
                print(f"⚠️  Файл {inner_image} не найден")

    doc.save(output_file)
    print(f"📄 Отчёт сохранён: {output_file}")
